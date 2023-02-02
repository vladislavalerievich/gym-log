from django.db.models import Q, Max
from django.http import FileResponse, HttpResponse
from rest_framework import viewsets, status
from rest_framework.decorators import action
from rest_framework.permissions import IsAuthenticatedOrReadOnly, IsAuthenticated
from rest_framework.response import Response
from io import BytesIO
import docx
from django.utils.text import slugify
from django.template.loader import render_to_string

from core.models import Exercise, Workout, Profile
from core.permissions import IsOwner
from core.serializers import ExerciseSerializer, WorkoutSerializer, ProfileSerializer

from core.models import WorkoutExercise, WorkoutExerciseDetail


class ProfileViewSet(viewsets.ModelViewSet):
    permission_classes = (IsAuthenticated, IsOwner)
    http_method_names = ['get', 'post']
    serializer_class = ProfileSerializer
    queryset = Profile.objects.all()

    def get_queryset(self):
        return self.queryset.filter(user=self.request.user.id)


class ExerciseViewSet(viewsets.ModelViewSet):
    permission_classes = (IsAuthenticatedOrReadOnly,)
    http_method_names = ['get', 'post', 'delete']
    serializer_class = ExerciseSerializer
    queryset = Exercise.objects.all()

    def get_queryset(self):
        return self.queryset.filter(Q(user=self.request.user.id) | Q(user=None)).order_by('name')

    @action(detail=False, methods=['get'])
    def body_parts(self, request):
        return Response(Exercise.body_part_list, status=status.HTTP_200_OK)

    @action(detail=False, methods=['get'])
    def equipment(self, request):
        return Response(Exercise.equipment_list, status=status.HTTP_200_OK)

    @action(detail=False, methods=['get'])
    def doc(self, request):
        doc = docx.Document()

        # Add the header to the document
        doc.add_heading('Exercise List', 0)

        # Create a dictionary to store the body parts and their corresponding exercises
        body_part_exercises = {}

        # Loop through all exercises
        for exercise in self.get_queryset():
            # Get the body part of the exercise
            body_part = exercise.body_part

            # If the body part isn't in the dictionary, add it
            if body_part not in body_part_exercises:
                body_part_exercises[body_part] = []

            # Add the exercise to the list of exercises for the body part
            body_part_exercises[body_part].append(exercise)

            # Sort the dictionary alphabetically by body part
        sorted_body_part_exercises = sorted(body_part_exercises.items(), key=lambda x: x[0])

        # Loop through the dictionary of body parts and exercises
        for body_part, exercises in sorted_body_part_exercises:
            # Add a heading for the body part
            doc.add_heading(body_part, level=1)

            # Loop through the exercises for the body part
            for exercise in exercises:
                # Add a bullet point for the exercise name and equipment
                doc.add_paragraph(exercise.name + ' (' + exercise.equipment + ')', style='List Bullet')

                # Add a paragraph for the exercise description
                doc.add_paragraph(exercise.description)

        # Save the document to a memory buffer
        file = BytesIO()
        doc.save(file)

        # Seek to the start of the file
        file.seek(0)

        # Return the file as a response
        response = FileResponse(file, content_type='application/msword')
        response['Content-Disposition'] = 'attachment; filename=exercise_list.docx'

        return response


class WorkoutViewSet(viewsets.ModelViewSet):
    permission_classes = (IsAuthenticated, IsOwner,)
    http_method_names = ['get', 'post', 'delete']
    serializer_class = WorkoutSerializer
    queryset = Workout.objects.all().order_by('-created')

    def get_queryset(self):
        return self.queryset.filter(user=self.request.user.id)

    @action(detail=False, methods=['get'])
    def doc(self, request):
        workouts = self.queryset.filter(user=self.request.user.id)
        profile = Profile.objects.get(user=self.request.user)

        doc = docx.Document()
        # Add the header to the document
        doc.add_heading('Workout history', 0)
        doc.add_heading("Email: " + profile.user.email, 2)

        for workout in workouts:
            doc.add_heading('Workout', 1)
            doc.add_heading(f'Status: {workout.status} | {workout.created.strftime("%d-%m-%Y %H:%M")}\n',
                            level=2)
            workout_exercises = WorkoutExercise.objects.filter(workout=workout)

            doc.add_heading('Exercises', 2)
            for workout_exercise in workout_exercises:
                exercise = workout_exercise.exercise
                doc.add_paragraph(exercise.name + ' (' + exercise.equipment + ')')
                workout_exercise_details = WorkoutExerciseDetail.objects.filter(workout_exercise=workout_exercise)
                for workout_exercise_detail in workout_exercise_details:
                    doc.add_paragraph(
                        f'{workout_exercise_detail.sets} sets × {workout_exercise_detail.reps} reps × {workout_exercise_detail.weight} {profile.weight_system}',
                        style='List Bullet')

        # Save the document to a memory buffer
        file = BytesIO()
        doc.save(file)

        # Seek to the start of the file
        file.seek(0)

        # Return the file as a response
        response = FileResponse(file, content_type='application/msword')
        response['Content-Disposition'] = 'attachment; filename=workout_history.docx'

        return response

    @action(detail=False, methods=['get'])
    def pr_doc(self, request):
        profile = Profile.objects.get(user=self.request.user)

        unique_exercises = WorkoutExercise.objects.filter(workout__user=request.user.id).values(
            'exercise').distinct().annotate(max_weight=Max('workout_exercise_details__weight'))

        doc = docx.Document()
        # # Add the header to the document
        doc.add_heading('Personal Records', 0)
        doc.add_heading(f'Heaviest weights {profile.user.email} have ever lifted:', 1)

        for exercise in unique_exercises:
            if exercise['max_weight'] > 0:
                exercise_obj = Exercise.objects.get(id=exercise['exercise'])
                doc.add_paragraph(
                    f"{exercise_obj.name} ({exercise_obj.equipment}) × {exercise['max_weight']} {profile.weight_system}",
                    style='List Bullet')

        # Save the document to a memory buffer
        file = BytesIO()
        doc.save(file)

        # Seek to the start of the file
        file.seek(0)

        # Return the file as a response
        response = FileResponse(file, content_type='application/msword')
        response['Content-Disposition'] = 'attachment; filename=personal_records.docx'

        return response
